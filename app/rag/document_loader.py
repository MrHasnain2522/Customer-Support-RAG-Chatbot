"""
Document Loader - Fixed for Product Catalogs
Ensures numbered items (1-10) stay within their own chunks.
"""
import os
import re
from typing import List, Dict
from app.utils.logger import get_logger

logger = get_logger(__name__)

class DocumentLoader:
    """Load documents from various file formats with Product-Aware chunking"""
    
    def __init__(self, knowledge_base_path: str = None):
        self.knowledge_base_path = knowledge_base_path or os.getenv(
            'KNOWLEDGE_BASE_PATH', 
            'knowledge_base/documents'
        )
        self.supported_extensions = ['.txt', '.md', '.pdf', '.docx']
    
    def load_all_documents(self) -> List[Dict]:
        documents = []
        if not os.path.exists(self.knowledge_base_path):
            logger.warning(f"Path not found: {self.knowledge_base_path}")
            return documents
        
        for root, dirs, files in os.walk(self.knowledge_base_path):
            for file in files:
                file_path = os.path.join(root, file)
                ext = os.path.splitext(file)[1].lower()
                if ext in self.supported_extensions:
                    try:
                        doc = self.load_document(file_path)
                        if doc: documents.append(doc)
                    except Exception as e:
                        logger.error(f"Error loading {file_path}: {str(e)}")
        return documents
    
    def load_document(self, file_path: str) -> Dict:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.txt': return self._load_txt(file_path)
        elif ext == '.md': return self._load_markdown(file_path)
        elif ext == '.pdf': return self._load_pdf(file_path)
        elif ext == '.docx': return self._load_docx(file_path)
        return None

    def _load_txt(self, file_path: str) -> Dict:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return {'text': text, 'metadata': {'source': file_path, 'filename': os.path.basename(file_path), 'type': 'txt'}}

    def _load_markdown(self, file_path: str) -> Dict:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return {'text': text, 'metadata': {'source': file_path, 'filename': os.path.basename(file_path), 'type': 'markdown'}}

    def _load_pdf(self, file_path: str) -> Dict:
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n\n"
            return {'text': text.strip(), 'metadata': {'source': file_path, 'filename': os.path.basename(file_path), 'type': 'pdf', 'pages': len(reader.pages)}}
        except Exception as e:
            logger.error(f"PDF Error: {str(e)}")
            return None

    def _load_docx(self, file_path: str) -> Dict:
        try:
            from docx import Document
            doc = Document(file_path)
            text = "\n\n".join([para.text for para in doc.paragraphs if para.text])
            return {'text': text, 'metadata': {'source': file_path, 'filename': os.path.basename(file_path), 'type': 'docx', 'paragraphs': len(doc.paragraphs)}}
        except Exception as e:
            logger.error(f"DOCX Error: {str(e)}")
            return None

    def chunk_text(self, text: str, chunk_size: int = 700, chunk_overlap: int = 100) -> List[str]:
        """
        Split text into chunks - IMPROVED FOR CATALOGS
        Maintains suit context by looking for product numbering (e.g., '1. ', '2. ')
        """
        if not text or len(text.strip()) == 0:
            return []
        
        # 1. First, try to split by Product Headings (e.g., "1. ", "2. ")
        # This keeps all info for one suit inside one chunk if possible
        product_split_pattern = r'\n(?=\d+\.\s+[A-Z\s]{5,})' 
        parts = re.split(product_split_pattern, text)
        
        final_chunks = []
        
        for part in parts:
            part = part.strip()
            if not part: continue
            
            # 2. If a single product section is too big for the chunk_size, 
            # fall back to the standard sliding window logic
            if len(part) > chunk_size:
                start = 0
                while start < len(part):
                    end = min(start + chunk_size, len(part))
                    # Try to find a good break point (newline or period)
                    if end < len(part):
                        for sep in ['\n\n', '\n', '. ']:
                            pos = part.rfind(sep, start, end)
                            if pos != -1 and pos > start + (chunk_size // 2):
                                end = pos + len(sep)
                                break
                    
                    chunk = part[start:end].strip()
                    if len(chunk) > 10:
                        final_chunks.append(chunk)
                    start = end - chunk_overlap
                    if start <= 0 or end >= len(part): break
            else:
                final_chunks.append(part)

        return final_chunks[:1000] # Safety limit

    def load_and_chunk_documents(self, chunk_size: int = 700, chunk_overlap: int = 100) -> List[Dict]:
        """
        Main entry point for the retriever.
        Returns the exact list of dictionaries your system expects.
        """
        documents = self.load_all_documents()
        chunked_docs = []
        
        for doc in documents:
            try:
                chunks = self.chunk_text(doc['text'], chunk_size, chunk_overlap)
                for i, chunk in enumerate(chunks):
                    chunked_docs.append({
                        'text': chunk,
                        'metadata': {
                            **doc['metadata'],
                            'chunk_index': i,
                            'total_chunks': len(chunks)
                        }
                    })
            except Exception as e:
                logger.error(f"Chunking error: {str(e)}")
        
        return chunked_docs